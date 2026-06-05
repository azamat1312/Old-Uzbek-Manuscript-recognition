"""Tanib olingan matnni fayllarga eksport qilish: TXT, DOCX, PDF.

Matn arab-fors yozuvida (o'ngdan-chapga, RTL) bo'lgani uchun har bir format
RTL ni to'g'ri qayta ishlaydi:

* TXT  — oddiy UTF-8 (yo'nalish matn muharririga bog'liq).
* DOCX — paragraf/run darajasida bidi/rtl bayroqlari qo'yiladi; shakllantirishni
         (harf bog'lanishini) Word o'zi bajaradi.
* PDF  — reportlab arab yozuvini o'zi shakllantirmaydi, shuning uchun matn
         `arabic_reshaper` bilan shakllantirilib, `python-bidi` bilan vizual
         tartibga keltiriladi va o'ngga tekislanib chiziladi.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from .config import PROJECT_ROOT

log = logging.getLogger("euyhtr.exporters")

# Loyiha bilan birga keladigan arab shrifti (tizimga bog'liq emas).
_BUNDLED_ARABIC_FONT = PROJECT_ROOT / "assets" / "fonts" / "Amiri-Regular.ttf"


# --------------------------------------------------------------------------- #
# TXT
# --------------------------------------------------------------------------- #
def export_txt(text: str, path: Path, title: Optional[str] = None) -> Path:
    header = f"{title}\n{'=' * len(title)}\n\n" if title else ""
    path.write_text(header + text, encoding="utf-8")
    return path


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def export_docx(text: str, path: Path, title: Optional[str] = None) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    ARABIC_FONT = "Arial"  # Windows'da arab yozuvini qo'llab-quvvatlaydi

    doc = Document()

    # Normal uslubni arab yozuvi uchun moslaymiz
    normal = doc.styles["Normal"]
    normal.font.name = ARABIC_FONT
    normal.font.size = Pt(14)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:cs"), ARABIC_FONT)  # complex script shrifti

    def _add_rtl_paragraph(line: str, *, bold: bool = False, size: int = 14):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement("w:bidi"))  # paragrafni RTL qilamiz
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)
        r_rpr = run._element.get_or_add_rPr()
        r_rpr.append(OxmlElement("w:rtl"))  # runni RTL qilamiz
        r_fonts = r_rpr.get_or_add_rFonts()
        r_fonts.set(qn("w:cs"), ARABIC_FONT)
        return p

    if title:
        _add_rtl_paragraph(title, bold=True, size=18)
        doc.add_paragraph()

    for line in text.split("\n"):
        _add_rtl_paragraph(line if line else "")

    doc.save(str(path))
    return path


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
# Arab yozuvini qo'llab-quvvatlovchi shriftlar (afzallik tartibida, bir nechta OT).
# Tizim shriftlari topilmasa, loyiha ichidagi Amiri shrifti (AmiriBundled) ishlatiladi.
_FONT_CANDIDATES = [
    # Loyiha ichidagi shrift — birinchi navbatda (tizimga bog'liq emas)
    ("AmiriBundled", str(_BUNDLED_ARABIC_FONT)),
    # Windows
    ("ArabicTypesetting", r"C:\Windows\Fonts\ARABTYPE.TTF"),
    ("SakkalMajalla", r"C:\Windows\Fonts\majalla.ttf"),
    ("Tahoma", r"C:\Windows\Fonts\tahoma.ttf"),
    ("Arial", r"C:\Windows\Fonts\arial.ttf"),
    ("ArialUnicode", r"C:\Windows\Fonts\ARIALUNI.TTF"),
    # Linux (Ubuntu/Debian)
    ("Amiri", "/usr/share/fonts/truetype/amiri/Amiri-Regular.ttf"),
    ("Amiri", "/usr/share/fonts/truetype/amiri/amiri-regular.ttf"),
    ("NotoNaskhArabic", "/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf"),
    ("KacstOne", "/usr/share/fonts/truetype/kacst-one/KacstOne.ttf"),
    ("KacstBook", "/usr/share/fonts/truetype/kacst/KacstBook.ttf"),
    ("FreeSerif", "/usr/share/fonts/truetype/freefont/FreeSerif.ttf"),
    # macOS
    ("GeezaPro", "/System/Library/Fonts/Supplemental/GeezaPro.ttf"),
]

_registered_font: Optional[str] = None


def _arabic_font_via_fontconfig() -> Optional[tuple]:
    """Linux: fontconfig (fc-match) orqali arab tilini qo'llaydigan shrift faylini topadi.
    Ma'lum yo'llar topilmaganda oxirgi chora sifatida ishlatiladi."""
    import shutil
    import subprocess

    if not shutil.which("fc-match"):
        return None
    try:
        out = subprocess.run(
            ["fc-match", "-f", "%{file}", ":lang=ar"],
            capture_output=True, text=True, timeout=5, check=False,
        )
        path = out.stdout.strip()
        if path and Path(path).exists():
            return ("SystemArabic", path)
    except Exception:  # noqa: BLE001
        pass
    return None


def _register_arabic_font() -> Optional[str]:
    """Mavjud bo'lgan birinchi arab shriftini reportlab'ga ro'yxatdan o'tkazadi.
    Avval ma'lum yo'llar, so'ng (Linux'da) fontconfig orqali qidiriladi.
    Topilmasa None qaytaradi (standart Helvetica ishlatiladi — arabcha ko'rinmaydi)."""
    global _registered_font
    if _registered_font is not None:
        return _registered_font

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = list(_FONT_CANDIDATES)
    fc = _arabic_font_via_fontconfig()
    if fc is not None:
        candidates.append(fc)

    for name, font_path in candidates:
        if Path(font_path).exists():
            try:
                pdfmetrics.registerFont(TTFont(name, font_path))
                _registered_font = name
                log.info("PDF uchun shrift: %s (%s)", name, font_path)
                return name
            except Exception as exc:  # noqa: BLE001
                log.warning("Shriftni yuklab bo'lmadi %s: %s", font_path, exc)
    log.warning("Arab shrifti topilmadi — PDF'da arabcha matn noto'g'ri ko'rinishi mumkin.")
    return None


def _shape_rtl(line: str) -> str:
    """Arab matnini PDF uchun shakllantirib, vizual (RTL) tartibga keltiradi."""
    if not line.strip():
        return ""
    import arabic_reshaper
    from bidi.algorithm import get_display

    return get_display(arabic_reshaper.reshape(line))


def export_pdf(text: str, path: Path, title: Optional[str] = None) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas

    font_name = _register_arabic_font() or "Helvetica"
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    margin = 2 * cm
    line_height = 22
    right_x = width - margin
    y = height - margin

    if title:
        c.setFont(font_name, 18)
        c.drawRightString(right_x, y, _shape_rtl(title))
        y -= line_height * 1.6

    c.setFont(font_name, 14)
    for raw in text.split("\n"):
        if y < margin:
            c.showPage()
            c.setFont(font_name, 14)
            y = height - margin
        c.drawRightString(right_x, y, _shape_rtl(raw))
        y -= line_height

    c.save()
    return path


# --------------------------------------------------------------------------- #
# Dispatcher
# --------------------------------------------------------------------------- #
EXPORTERS = {
    "txt": export_txt,
    "docx": export_docx,
    "pdf": export_pdf,
}

MEDIA_TYPES = {
    "txt": "text/plain; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


def export(fmt: str, text: str, path: Path, title: Optional[str] = None) -> Path:
    if fmt not in EXPORTERS:
        raise ValueError(f"Noma'lum format: {fmt}")
    return EXPORTERS[fmt](text, path, title)
